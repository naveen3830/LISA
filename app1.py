import streamlit as st
import pandas as pd
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate, PromptTemplate
from langchain.prompts import SystemMessagePromptTemplate, HumanMessagePromptTemplate
from langchain_core.messages import HumanMessage
from langchain.schema import AIMessage
from langchain_core.output_parsers import StrOutputParser
from pandasql import sqldf
from functions import check
from dotenv import load_dotenv
from pathlib import Path
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
import re

st.set_page_config(page_title="LISA : LLM Informed Statistical Analysis", page_icon=":books:", layout="wide")

if 'df' not in st.session_state:
    st.session_state.df = None
if 'results' not in st.session_state:
    st.session_state.results = {}
if 'tables' not in st.session_state:
    st.session_state.tables = {}

tab1, tab2, tab3, tab4 = st.tabs(["Home", "ChatBot", "Report Generation", "About"])

def get_llm_response(llm, prompt_template, data):
    system_message_prompt = SystemMessagePromptTemplate.from_template(
        "You are StatBot, an expert statistical analyst. "
        "Explain the output in simple English. Straight away start with your explanations.")
    
    human_message_prompt = HumanMessagePromptTemplate.from_template(prompt_template)
    
    chat_prompt = ChatPromptTemplate.from_messages([system_message_prompt, human_message_prompt])
    formatted_chat_prompt = chat_prompt.format_messages(**data)
    response = llm.invoke(formatted_chat_prompt)
    return response.content

def groq_infer(llm, prompt):
    messages = [HumanMessage(content=prompt)]
    response = llm(messages)
    print(response.content)
    return response.content

def generate_explanation_prompt(uploaded_file, question, result):
    return f"""
    Given the context of the dataset from {uploaded_file}, provide a thoughtful explanation of the following answer in simple English. Consider the original question and interpret the results in that context. Do not explain the SQL query itself.

    Original question: {question}

    Result:
    {result}

    Please explain:
    1. What does this result mean in relation to the question asked?
    2. What insights can we draw from this data?
    3. Are there any notable patterns or anomalies in the result?
    4. How might this information be useful or actionable?

    Provide your explanation in a clear, concise manner that a non-technical person could understand.
    """
    
def add_table_to_doc(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
    for row in df.itertuples():
        row_cells = table.add_row().cells
        for i, value in enumerate(row[1:], start=0):
            row_cells[i].text = str(value)

def format_markdown(text):
    # Remove asterisks for bullet points
    text = re.sub(r'^\s*\*\s*', '• ', text, flags=re.MULTILINE)
    
    # Convert numbered lists to proper format
    text = re.sub(r'^\s*(\d+)\.\s*', lambda m: f"{m.group(1)}. ", text, flags=re.MULTILINE)
    
    return text

def add_formatted_text_to_doc(doc, text):
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip().startswith('•'):
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.text = para.strip()[2:]  # Remove the bullet point
        elif re.match(r'^\d+\.', para.strip()):
            p = doc.add_paragraph()
            p.style = 'List Number'
            p.text = para.strip()
        else:
            doc.add_paragraph(para)

with st.sidebar:
    with st.sidebar.expander("Get Your API Key Here"):
        st.markdown("## How to use\n"
            "1. Enter your [Groq API key](https://console.groq.com/keys) below🔑\n" 
            "2. Upload a CSV file📄\n"
            "3. Let LISA do its work!!!💬\n")
    
    groq_api_key = st.text_input("Enter your Groq API key:", type="password",
            placeholder="Paste your Groq API key here (gsk_...)",
            help="You can get your API key from https://console.groq.com/keys")
    
    with st.sidebar.expander("Model Parameters"):
        model_name = st.selectbox("Select Model:",["llama-3.1-70b-versatile","llama3-70b-8192", "mixtral-8x7b-32768", "gemma2-9b-it"])
        temperature = st.slider("Temperature: It determines whether the output is more random, creative or more predictable.", min_value=0.0, max_value=1.0, value=0.5, step=0.1)
        top_p = st.slider("Top-p: It determines the cumulative probability distribution used for sampling the next token in the generated response", min_value=0.0, max_value=1.0, value=1.0, step=0.25)

    st.divider()

    llm = None
    if groq_api_key:
        try:
            llm = ChatGroq(
                groq_api_key=groq_api_key, 
                model_name=model_name,
                temperature=temperature,
                top_p=top_p
            )
        except Exception as e:
            st.sidebar.error(f"Error initializing model: {str(e)}")

template = """You are a powerful text-to-SQL model. Your job is to answer questions about a database. You are given a question and context regarding one or more tables. Don't add \n characters.

You must output the SQL query that answers the question in a single line.

### Input:
`{question}`

### Context:
`{context}`

### Response:
"""

prompt = PromptTemplate.from_template(template=template)

with tab1:
    st.header("Welcome to LISA: LLM Informed Statistical Analysis 🎈", divider='rainbow')
    st.markdown("LISA is an innovative platform designed to automate your data analysis process using advanced Large Language Models (LLM) for insightful inferences. Whether you're a data enthusiast, researcher, or business analyst, LISA simplifies complex data tasks, providing clear and comprehensible explanations for your data.")
    st.markdown("LISA combines the efficiency of automated data processing with the intelligence of modern language models to deliver a seamless and insightful data analysis experience. Empower your data with LISA!")
    st.divider()
    
    uploaded_file = st.file_uploader("Upload a CSV file", type=['csv'])
    
    if uploaded_file is not None:
        st.session_state.df = pd.read_csv(uploaded_file, encoding="latin1")
        st.session_state.df.columns = st.session_state.df.columns.str.replace(r"[^a-zA-Z0-9_]", "", regex=True)
        st.dataframe(st.session_state.df)
        st.divider()

        option = st.selectbox("Select an option:", ["Show dataset dimensions", "Display data description", "Verify data integrity", "Summarize numerical data statistics", "Summarize categorical data", "Ask a question about the data"])
        
        if not groq_api_key:
            st.warning("Please enter your Groq API key in the sidebar to use the analysis features.")
        elif llm is None:
            st.error("Failed to initialize the model. Please check your API key.")
        else:
            if option == "Show dataset dimensions":
                shape_of_the_data = st.session_state.df.shape
                response = get_llm_response(llm, 'The shape of the dataset is: {shape}', {'shape': shape_of_the_data})
                st.write(response)
                st.session_state.results['Dataset Dimensions'] = response
                st.session_state.tables['Dataset Dimensions'] = pd.DataFrame({'Rows': [shape_of_the_data[0]], 'Columns': [shape_of_the_data[1]]})
                
            elif option == "Display data description":
                column_description = st.session_state.df.columns
                response = get_llm_response(llm, 'The columns in the dataset are: {columns}. Provide a brief description for each column.', {'columns': column_description})
                st.write(response)
                st.session_state.results['Data Description'] = response
                st.session_state.tables['Data Description'] = pd.DataFrame({'Columns': column_description})
                
            elif option == "Verify data integrity":
                df_check = check(st.session_state.df)
                st.dataframe(df_check)
                st.divider()
                response = get_llm_response(llm, 'The data integrity check results are: {df_check}', {'df_check': df_check})
                st.write(response)
                st.session_state.results['Data Integrity'] = response
                st.session_state.tables['Data Integrity'] = df_check
                
            elif option == "Summarize numerical data statistics":
                describe_numerical = st.session_state.df.describe().T
                st.dataframe(describe_numerical)
                st.divider()    
                response = get_llm_response(llm, 'The numerical data statistics are: {stats}', {'stats': describe_numerical})
                st.write(response)
                st.session_state.results['Numerical Data Statistics'] = response
                st.session_state.tables['Numerical Data Statistics'] = describe_numerical
                
            elif option == "Summarize categorical data":
                categorical_df = st.session_state.df.select_dtypes(include=['object'])
                if categorical_df.empty:
                    st.write("No categorical columns found.")
                    response = get_llm_response(llm, 'There are no categorical columns in this dataset.', {})
                else:
                    describe_categorical = categorical_df.describe()
                    st.dataframe(describe_categorical)
                    st.divider()
                    response = get_llm_response(llm, 'The categorical data summary is: {summary}', {'summary': describe_categorical})
                st.write(response)
                st.session_state.results['Categorical Data Summary'] = response
                if not categorical_df.empty:
                    st.session_state.tables['Categorical Data Summary'] = describe_categorical
            
            elif option == "Ask a question about the data":
                question = st.text_input("Write a question about the data", key="question")
                if question:
                    attempt = 0
                    max_attempts = 5
                    while attempt < max_attempts:
                        try:
                            context = pd.io.sql.get_schema(st.session_state.df.reset_index(), "df").replace('"', "")
                            input_data = {"context": context, "question": question}
                            formatted_prompt = prompt.format(**input_data)
                            response = groq_infer(llm, formatted_prompt)
                            final = response.replace("`", "").replace("sql", "").strip()
                            
                            result = sqldf(final, {'df': st.session_state.df})
                            st.write("Answer:")
                            st.dataframe(result)
                            
                            explanation_prompt = generate_explanation_prompt(uploaded_file.name, question, result.to_string())
                            explanation_response = groq_infer(llm, explanation_prompt)
                            st.write("Explanation:")
                            st.write(explanation_response)
                            st.session_state.results['Custom Question'] = f"Question: {question}\nAnswer: {explanation_response}"
                            st.session_state.tables['Custom Question'] = result
                            break
                        except Exception as e:
                            attempt += 1
                            st.error(f"Attempt {attempt}/{max_attempts} failed. Error: {str(e)}")
                            if attempt == max_attempts:
                                st.error("Unable to get the correct query after 5 attempts. Please try again or refine your question.")
                            continue
                else:
                    st.warning("Please enter a question before clicking 'Get Answer'.")
                    
with tab2:
    st.markdown("""Our integrated chatbot is available to assist you, providing real-time answers to your data-related queries and enhancing your overall experience with personalized support.""")
    st.markdown("""---""")

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    def get_response(query, chat_history, df):
        template = """
        You are a knowledgeable data assistant. Answer the user's question based on the provided dataset and the conversation history. If the data isn't directly related to the question, guide the user on how they might extract relevant insights also do not directly provide the code for every question provide code only when user asks for.

        Dataset information (limit to first 50 rows):
        {df}

        Chat history:
        {chat_history}

        User question: {user_question}

        If the question is not clear, ask for clarification. Provide SQL queries or Python code snippets where appropriate to help the user interact with their data.
        """
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | llm | StrOutputParser()
        return chain.stream({
            "chat_history": chat_history,
            "user_question": query,
            "df": df.head(50).to_string() if df is not None else "No data uploaded yet."
        })

    # Display chat history
    for message in st.session_state.chat_history:
        if isinstance(message, HumanMessage):
            with st.chat_message("Human"):
                st.markdown(message.content)
        else:
            with st.chat_message("AI"):
                st.markdown(message.content)
                
    if not groq_api_key:
        st.warning("Please enter your Groq API key in the sidebar to use the chatbot.")
    elif llm is None:
        st.error("Failed to initialize the model. Please check your API key.")
    else:
        st.write("") 
        user_query = st.chat_input("Type your message here")
        
        if user_query is not None and user_query != "":
            st.session_state.chat_history.append(HumanMessage(content=user_query))
            
            with st.chat_message("Human"):
                st.markdown(user_query)
                
            with st.chat_message("AI"):
                full_response = ""
                message_placeholder = st.empty()
                try:
                    for chunk in get_response(user_query, st.session_state.chat_history, st.session_state.get('df')):
                        full_response += chunk
                        message_placeholder.markdown(full_response + "")
                    message_placeholder.markdown(full_response)
                except Exception as e:
                    error_message = f"An error occurred: {str(e)}. Please make sure you've uploaded a dataset."
                    message_placeholder.error(error_message)
                    full_response = error_message
            
            st.session_state.chat_history.append(AIMessage(content=full_response))

with tab3:
    st.header("Report Generation")
    
    if st.session_state.results:
        st.write("Select the sections you want to include in your report:")
        
        selected_sections = {}
        for key, value in st.session_state.results.items():
            selected_sections[key] = st.checkbox(key, value=True)
        
        if st.button("Generate Report"):
            doc = Document()
            
            # Modify existing Normal style
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Add List Bullet style if it doesn't exist
            if 'List Bullet' not in doc.styles:
                list_bullet_style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
                list_bullet_style.base_style = doc.styles['Normal']
                list_bullet_style.paragraph_format.left_indent = Inches(0.25)
                list_bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
            
            # Add List Number style if it doesn't exist
            if 'List Number' not in doc.styles:
                list_number_style = doc.styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
                list_number_style.base_style = doc.styles['Normal']
                list_number_style.paragraph_format.left_indent = Inches(0.25)
                list_number_style.paragraph_format.first_line_indent = Inches(-0.25)
            
            doc.add_heading('LISA: LLM Informed Statistical Analysis Report', 0)
            
            for key, selected in selected_sections.items():
                if selected:
                    doc.add_heading(key, level=1)
                    formatted_text = format_markdown(st.session_state.results[key])
                    add_formatted_text_to_doc(doc, formatted_text)
                    if key in st.session_state.tables:
                        doc.add_paragraph("Table output:")
                        add_table_to_doc(doc, st.session_state.tables[key])
            
            bio = BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="Download Report",
                data=bio.getvalue(),
                file_name="LISA_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.write("No results available. Please analyze data in the Home tab first.")
        
        
with tab4:
    st.header("LLM Model Card",divider="rainbow")
    
    st.markdown("In our innovative project LISA (LLM Informed Statistical Analysis), we are harnessing the power of Groq-hosted large language models (LLMs) to revolutionize the way statistical analysis is performed and interpreted. Groq’s platform plays a pivotal role in enabling LISA to deliver accurate, fast, and insightful data analysis by providing access to highly optimized, open-source LLMs that are tailored for complex data processing tasks.")
    
    st.markdown("Groq is the AI infrastructure company that delivers fast AI inference.The LPU™ Inference Engine by Groq is a hardware and software platform that delivers exceptional compute speed, quality, and energy efficiency.")
    
    st.markdown("The table below provides comparision of the performance of different LLM models across various NLP (Natural Language Processing) benchmarks")
    
    data_folder_path = "Data"
    dataframes = {}
    for file_name in os.listdir(data_folder_path):
        if file_name.endswith(".csv"):
            data_file_path = os.path.join(data_folder_path, file_name)
            df = pd.read_csv(data_file_path)
            dataframes[file_name] = df

    model_card = dataframes.get('modelcard.csv')
    if model_card is not None:
        st.dataframe(model_card, hide_index=True)
    else:
        st.error("Model card CSV not found.")
        
    st.markdown("""
<style>
ul {
    list-style-type: disc;
    margin-left: 20px;
}
</style>
<ul>
    Here’s what these benchmarks mean:
    <li><b>MMLU (Massive Multitask Language Understanding):</b> A benchmark designed to understand how well a language model can multitask. The model’s performance is assessed across a range of subjects, such as math, computer science, and law.</li>
    <li><b>GPQA (Graduate-Level Google-Proof Q&A):</b> Assesses a model’s ability to answer questions that are challenging for search engines to solve directly. This benchmark evaluates whether the AI can handle questions that usually require human-level research skills.</li>
    <li><b>HumanEval:</b> Assesses how well the model can write code by asking it to perform programming tasks.</li>
    <li><b>GSM-8K:</b> Evaluates the model’s ability to solve math word problems.</li>
    <li><b>MATH:</b> Tests the model’s ability to solve middle school and high school math problems.</li>
</ul>
""", unsafe_allow_html=True)
    
    st.info("We've observed that the Gemma2-9B-IT model excels in querying data, while the Llama variants are particularly effective for inferring results.", icon="ℹ️")